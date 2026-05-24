#!/usr/bin/env python3
"""
Complete extraction: content, images, with proper paragraph structure.

Fixes:
1. Preserve paragraph breaks (don't merge everything into one string)
2. Extract images and map them to their paragraph locations
3. Correct task/content separation
"""

import json
import re
import os
import zipfile
import shutil
from docx import Document
from lxml import etree

DOCX_PATH = '/home/user/project/book(t)/новый материал.docx'
APP_DIR = '/home/user/project/book(t)/app'
IMG_DIR = os.path.join(APP_DIR, 'images')

# ============================================================
# Step 1: Extract images
# ============================================================
os.makedirs(IMG_DIR, exist_ok=True)

image_map = {}  # rId -> filename

with zipfile.ZipFile(DOCX_PATH) as z:
    # Find all image files
    for name in z.namelist():
        if name.startswith('word/media/'):
            basename = os.path.basename(name)
            # Extract to app/images/
            target = os.path.join(IMG_DIR, basename)
            with z.open(name) as src:
                with open(target, 'wb') as dst:
                    dst.write(src.read())
            image_map[basename] = basename

print(f"Extracted {len(image_map)} images to {IMG_DIR}")

# ============================================================
# Step 2: Map image rIds to paragraph positions in the document
# ============================================================
doc = Document(DOCX_PATH)

# Build rId -> image filename mapping
rid_to_image = {}
for rel_id, rel in doc.part.rels.items():
    if 'image' in rel.reltype:
        target = rel.target_ref
        if target.startswith('media/'):
            fname = os.path.basename(target)
            rid_to_image[rel_id] = fname

# Find images in paragraphs
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'v': 'urn:schemas-microsoft-com:vml',
}

para_images = {}  # paragraph_index -> [image_filenames]

for i, p in enumerate(doc.paragraphs):
    images = []

    # Check for drawings (modern format)
    for blip in p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if embed and embed in rid_to_image:
            images.append(rid_to_image[embed])

    # Check for inline shapes (older format)
    for imagedata in p._element.findall('.//{urn:schemas-microsoft-com:vml}imagedata'):
        rid = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        if rid and rid in rid_to_image:
            images.append(rid_to_image[rid])

    if images:
        para_images[i] = images

print(f"Found images in {len(para_images)} paragraphs")

# Check for paragraphs that only have images (no text) - they're standalone figures
standalone_images = {}
for idx, imgs in para_images.items():
    text = doc.paragraphs[idx].text.strip()
    if not text or len(text) < 10:
        standalone_images[idx] = imgs

# ============================================================
# Step 3: Parse document structure
# ============================================================
def clean_text(text):
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text

paragraphs = []
for i, p in enumerate(doc.paragraphs):
    text = clean_text(p.text)
    imgs = para_images.get(i, [])
    # Check if this paragraph has an image but no meaningful text
    if imgs and (not text or len(text) < 5):
        text = None  # Will be handled as image-only

    paragraphs.append({
        'index': i,
        'text': text,
        'style': p.style.name,
        'images': imgs
    })

# State machine
modules = []
current_module = None
current_lecture = None
current_lab = None

ST_CONTENT = 'content'
ST_QUESTIONS = 'questions'
ST_REFERENCES = 'references'
ST_PRACTICAL = 'practical'
ST_TEST = 'test'

current_state = ST_CONTENT

def is_module(text):
    return bool(re.match(r'^Модуль\s+\d+', text)) if text else False

def is_lecture(text):
    return bool(re.match(r'^\d+\s*[-–]\s*дәріс[.\s]', text, re.IGNORECASE)) if text else False

def is_lab(text):
    return bool(re.match(r'^№\s*\d+[\s-]*Зертханалық\s+сабақ', text)) if text else False

def is_practical_tasks_header(text):
    return bool(re.match(r'^Практикалық\s+тапсырмалар', text)) if text else False

def is_practice_plan(text):
    return 'Практика жоспары' in text if text else False

def is_control_questions(text):
    return 'Бақылау сұрақтары' in text if text else False

def is_references(text):
    return 'Пайдаланылған әдебиеттер' in text if text else False

def is_variant_header(text):
    return bool(re.match(r'^\d+\s*[-–]\s*нұсқа', text, re.IGNORECASE)) if text else False

def is_task_like(text):
    return bool(re.match(
        r'^(Тапсырма\s*\d|Жаттығу[:.]|Тапсырманың\s+берілгені|'
        r'Тапсырма[:.]|Тапсырмалар[:]|'
        r'№\s*\d+\s*тапсырма|'
        r'\d+[\s-]*тапсырма[.\s])',
        text, re.IGNORECASE
    )) if text else False

def is_section_label(text):
    if not text:
        return False
    labels = ['Жұмыс барысы', 'Нұсқау:', 'Нұсқаулық:', 'Нұсқау', 'Нұсқаулық',
              'Мақсаты:', 'Тақырыбы:']
    return any(text.strip().startswith(l) for l in labels)

def make_para(text, imgs):
    """Create a paragraph object with text and optional images."""
    obj = {}
    if text:
        obj['text'] = text
    if imgs:
        obj['images'] = imgs
    return obj

# Process all paragraphs
for p in paragraphs:
    text = p['text']
    imgs = p['images']
    style = p['style']

    # ---- Module ----
    if text and is_module(text):
        if current_module:
            modules.append(current_module)
        current_module = {
            'title': text,
            'lectures': [],
            'labs': [],
            'intro': None
        }
        current_lecture = None
        current_lab = None
        current_state = ST_CONTENT
        continue

    # ---- Lecture ----
    if text and is_lecture(text):
        if current_module is None:
            current_module = {'title': 'АКТ пәні', 'lectures': [], 'labs': [], 'intro': None}
        current_lecture = {
            'title': text,
            'paragraphs': [],
            'control_questions': [],
            'references': []
        }
        current_lab = None
        current_module['lectures'].append(current_lecture)
        current_state = ST_CONTENT
        continue

    # ---- Lab ----
    if text and is_lab(text):
        if current_module is None:
            current_module = {'title': 'АКТ пәні', 'lectures': [], 'labs': [], 'intro': None}
        current_lab = {
            'title': text,
            'paragraphs': [],
            'practical_tasks': [],
            'test_tasks': [],
            'control_questions': [],
            'references': []
        }
        current_lecture = None
        current_module['labs'].append(current_lab)
        current_state = ST_CONTENT
        continue

    # ---- State transitions ----
    target = current_lecture or current_lab

    if text and is_practical_tasks_header(text):
        current_state = ST_PRACTICAL
        continue

    if text and is_practice_plan(text):
        current_state = ST_CONTENT
        if target:
            target.setdefault('paragraphs', []).append(make_para(text, imgs))
        continue

    if text and is_control_questions(text):
        current_state = ST_QUESTIONS
        continue

    if text and is_references(text):
        current_state = ST_REFERENCES
        if target:
            target.setdefault('references', []).append(text)
        continue

    if text and 'тест тапсырма' in text.lower() and ('құр' in text.lower() or 'жаса' in text.lower()):
        current_state = ST_TEST
        if target:
            target.setdefault('test_tasks', []).append(text)
        continue

    # Skip truly empty paragraphs (no text AND no images)
    if not text and not imgs:
        continue

    # ---- Route ----
    if target:
        if current_state == ST_QUESTIONS:
            if text:
                target.setdefault('control_questions', []).append(text)
        elif current_state == ST_REFERENCES:
            if text:
                target.setdefault('references', []).append(text)
        elif current_state == ST_PRACTICAL:
            target.setdefault('practical_tasks', []).append(make_para(text, imgs))
        elif current_state == ST_TEST:
            if text:
                target.setdefault('test_tasks', []).append(text)
        else:
            target.setdefault('paragraphs', []).append(make_para(text, imgs))
    elif current_module is not None and text:
        if current_module['intro'] is None:
            current_module['intro'] = {'paragraphs': []}
        current_module['intro']['paragraphs'].append(make_para(text, imgs))
    elif current_module is not None and imgs:
        # Image-only paragraph in intro
        if current_module['intro'] is None:
            current_module['intro'] = {'paragraphs': []}
        current_module['intro']['paragraphs'].append(make_para(None, imgs))

if current_module:
    modules.append(current_module)

# ============================================================
# Step 4: Post-process practical_tasks - group by task header
# ============================================================
for mod in modules:
    for lab in mod.get('labs', []):
        tasks = lab.get('practical_tasks', [])
        if not tasks:
            # Try to extract tasks embedded in paragraphs
            paras = lab.get('paragraphs', [])
            if paras:
                new_paras = []
                new_tasks = []
                task_buf = []
                in_task = False

                for para in paras:
                    t = para.get('text', '')
                    if t and is_task_like(t):
                        if task_buf:
                            new_tasks.append(task_buf)
                        task_buf = [para]
                        in_task = True
                    elif in_task:
                        task_buf.append(para)
                    else:
                        new_paras.append(para)

                if task_buf:
                    new_tasks.append(task_buf)

                if len(new_tasks) >= 2:
                    lab['paragraphs'] = new_paras
                    lab['practical_tasks'] = new_tasks
            continue

        # Group tasks by header
        grouped = []
        buf = []

        for para in tasks:
            t = para.get('text', '')
            if t and (is_task_like(t) or is_variant_header(t)):
                if buf:
                    grouped.append(buf)
                buf = [para]
            elif t and is_section_label(t) and not buf:
                # This is a label before tasks → move to content
                lab.setdefault('paragraphs', []).append(para)
                continue
            else:
                buf.append(para)

        if buf:
            grouped.append(buf)

        lab['practical_tasks'] = grouped

# ============================================================
# Step 5: Build output
# ============================================================
output = {
    'title': 'Ақпараттық-коммуникациялық технологиялар',
    'title_short': 'АКТ',
    'description': 'Пәннің оқу-әдістемелік кешені',
    'institution': 'М.Дулатов атындағы Қостанай инженерлік-экономикалық университеті',
    'department': 'Ақпараттық технологиялар және автоматика кафедрасы',
    'year': '2025',
    'language': 'kk',
    'modules': modules
}

with open(os.path.join(APP_DIR, 'content.json'), 'w', encoding='utf-8') as f:
    json.dump(output, f, ensure_ascii=False, indent=2)

# ============================================================
# Statistics
# ============================================================
print(f"\nModules: {len(output['modules'])}")
for i, m in enumerate(output['modules']):
    print(f"\n  Module {i}: {m['title']}")
    for j, lec in enumerate(m.get('lectures', [])):
        paras = lec.get('paragraphs', [])
        n_img = sum(1 for p in paras if p.get('images'))
        nq = len(lec.get('control_questions', []))
        print(f"    L{j}: {lec['title'][:100]} ({len(paras)}p, {n_img}img, {nq}q)")
    for j, lab in enumerate(m.get('labs', [])):
        paras = lab.get('paragraphs', [])
        n_img = sum(1 for p in paras if p.get('images'))
        tasks = lab.get('practical_tasks', [])
        total_task_paras = sum(len(t) for t in tasks)
        print(f"    Lab{j}: {lab['title'][:90]} ({len(paras)}p, {n_img}img, {len(tasks)} tasks, {total_task_paras} task-paras)")

total_lec = sum(len(m.get('lectures',[])) for m in output['modules'])
total_lab = sum(len(m.get('labs',[])) for m in output['modules'])
total_tasks = sum(len(lab.get('practical_tasks',[])) for m in output['modules'] for lab in m.get('labs',[]))
total_imgs = len(image_map)
print(f"\n=== TOTAL: {len(output['modules'])} modules, {total_lec} lectures, {total_lab} labs ===")
print(f"=== {total_tasks} task groups, {total_imgs} images ===")
